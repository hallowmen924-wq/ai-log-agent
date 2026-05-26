from docx import Document
from docx.shared import Pt


def add_heading(doc: Document, text: str, level: int = 1) -> None:
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(12 if level > 1 else 14)


def add_para(doc: Document, text: str) -> None:
    p = doc.add_paragraph(text)
    for run in p.runs:
        run.font.name = "Malgun Gothic"
        run.font.size = Pt(10.5)


def add_bullets(doc: Document, items: list[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        for run in p.runs:
            run.font.name = "Malgun Gothic"
            run.font.size = Pt(10.5)


def main() -> None:
    doc = Document()

    style = doc.styles["Normal"]
    style.font.name = "Malgun Gothic"
    style.font.size = Pt(10.5)

    add_heading(doc, "AI 에이전트 대회 결과보고서", level=0)
    add_para(doc, "에이전트명: Bunny 금융 워크스페이스 AI (Bunny Loan Copilot)")

    add_heading(doc, "1. 에이전트 개요", level=1)
    add_para(
        doc,
        "대출 심사/상담/상품개발/정책준수 업무를 하나의 워크스페이스에서 처리하는 멀티에이전트형 AI 시스템입니다. "
        "심사 로그·규제문서·뉴스·온톨로지 데이터를 통합하고, 질문 의도(Intent)에 따라 "
        "Explainability/Cluster/Policy/Creative 분석을 동적으로 수행합니다.",
    )

    add_heading(doc, "2. 당사 대출 심사 관련 현황 및 문제점", level=1)
    add_bullets(
        doc,
        [
            "심사 결과(거절사유, 금리, 한도)에 대한 상담원 설명 일관성 부족",
            "규제문서 근거 탐색이 수작업 중심이라 응대 속도 저하",
            "상품별/고객군별 인사이트가 분산되어 의사결정 리드타임 증가",
            "뉴스/외부 신호 반영 지연으로 리스크 선제 대응 한계",
            "이미지 심사결과(캡처) 기반 질의 응대 체계 부재",
        ],
    )

    add_heading(doc, "3. 개선사항", level=1)
    add_bullets(
        doc,
        [
            "Intent 기반 라우팅(승인요인/거절사유/금리·한도/정책·규제/창의질의)",
            "FAISS 중심 RAG 파이프라인 + 근거(Citation) 강화",
            "Policy/Ontology Reasoning 카드 우선 노출 로직 정교화",
            "OCR(이미지 드래그앤드롭) → 텍스트 추출 → 심사상담형 응답 자동화",
            "뉴스 수집 워커 주기/상태 가시화 및 장애 원인 분리(lock-busy/timeout/parse-fail)",
            "상품개발모드 토론 UX 경량화(단일 안건 선택, 단계 시각화, 중간상태 스트리밍)",
        ],
    )

    add_heading(doc, "4. 적용채널", level=1)
    add_bullets(
        doc,
        [
            "내부 상담원 데스크(웹)",
            "심사/기획/영업/IT 협업 워크스페이스",
            "정책/준법 검토 지원 화면",
        ],
    )

    add_heading(doc, "5. 활용데이터", level=1)
    add_bullets(
        doc,
        [
            "대출 심사 로그(data/logs)",
            "규제문서 PDF(업로드 문서)",
            "뉴스 수집 데이터(크롤링 + 요약/신호)",
            "온톨로지/공통피처/군집 산출물",
            "OCR 추출 텍스트(심사결과 이미지)",
        ],
    )

    add_heading(doc, "6. 데이터유형", level=1)
    add_bullets(
        doc,
        [
            "정형: 심사결과, 금리, 한도, 소득, 코드값",
            "반정형: JSON 메타데이터, 피처 매핑",
            "비정형: 규제문서 본문, 뉴스 본문, 이미지 텍스트(OCR)",
        ],
    )

    add_heading(doc, "7. 에이전트 구성", level=1)
    add_para(doc, "에이전트1: Explainability Agent (핵심 피처 및 거절/승인 요인 설명)")
    add_para(doc, "에이전트2: Customer Cluster Intelligence (고객군집별 승인률/금리/한도/리스크 비교)")
    add_para(doc, "에이전트3: Policy/Ontology Reasoning (규제 근거 탐색, 정책-답변 정합성 검증)")
    add_para(doc, "에이전트4: Product Development Agent (신상품/룰 보완 토론 및 실행안 요약)")

    add_heading(doc, "8. 오픈소스 구성", level=1)
    add_bullets(
        doc,
        [
            "LLM: Ollama 기반 로컬 LLM",
            "개발 F/W: FastAPI, React(Vite)",
            "벡터DB: FAISS",
            "데이터DB: 파일 기반 JSON/아티팩트(+ 선택적 Neo4j 확장)",
            "UI: React + Framer Motion + Cytoscape.js",
            "모델서빙: Ollama",
            "기타: RapidOCR/Paddle 계열 OCR, WebSocket",
        ],
    )
    add_para(
        doc,
        "기술셋 구성전략: 로컬 중심(보안/비용/지연시간) 구조를 우선 적용하고, "
        "API 기반 모듈화로 확장성 확보. FAISS의 검색 속도와 온톨로지 해석력을 결합해 실무 응답력을 강화.",
    )

    add_heading(doc, "9. RAG 구성 전략", level=1)
    add_bullets(
        doc,
        [
            "규제문서/뉴스/로그를 소스별 인덱싱",
            "Intent 기반 검색 우선순위 분기",
            "중복 제거·Top-K 제한·근거(Citation) 노출",
            "정책 질의는 Policy 카드 우선, 창의 질의는 경량 컨텍스트 + LLM 생성형 응답",
        ],
    )

    add_heading(doc, "10. 전처리", level=1)
    add_para(doc, "원본데이터: 심사 로그, 규제 PDF, 뉴스, 이미지")
    add_bullets(
        doc,
        [
            "문서 청킹 + 메타데이터(document/page/article/chunk_id) 보존",
            "코드/금리/한도/소득 등 핵심 필드 정규화",
            "OCR 텍스트에서 심사 핵심 엔티티 추출",
            "피처 공통화/군집 재산출 파이프라인 운영",
        ],
    )

    add_heading(doc, "11. 본처리", level=1)
    add_para(doc, "알고리즘: 의도분류, 벡터검색, 피처 중요도/군집비교/정책근거 매칭, LLM 요약·상담 멘트 생성")
    add_para(doc, "우선순위 전략: 정책·규제 질문은 Policy 우선, 심사설명은 Explainability 우선, 고객군 질문은 Cluster 우선")

    add_heading(doc, "12. AI 에이전트 구성도 - 시스템 구성도", level=1)
    add_para(
        doc,
        "사용자 UI → Intent Router → (Explainability/Cluster/Policy/Product Dev/OCR) → RAG Orchestrator → "
        "FAISS/온톨로지/뉴스데이터/Ollama → 결과 카드 출력",
    )

    add_heading(doc, "13. AI 에이전트 구성도 - 사용자 관점 업무 흐름도", level=1)
    add_para(
        doc,
        "질문/이미지 입력 → 의도 해석 → 질문 유형별 라우팅(정책/심사설명/군집/창의) → "
        "근거·분석 카드 생성 → 상담 멘트/실행안 제시",
    )

    add_heading(doc, "14. AI 에이전트 시연 (핵심 유스케이스 5개)", level=1)
    table = doc.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    header = table.rows[0].cells
    header[0].text = "유스케이스 명"
    header[1].text = "사용자 입력"
    header[2].text = "결과 출력"
    rows = [
        (
            "거절사유 상담 가이드",
            "이지론 거절 고객군에서 가장 자주 연결되는 reject reason은?",
            "Top 거절코드/비중 + 상담원용 원인→안내 멘트",
        ),
        (
            "승인요인 설명",
            "40대 직장인의 카드론 승인에 중요한 요인은?",
            "Top 3 피처 + 군집 보조 정보",
        ),
        (
            "규제근거 응답",
            "금리인하요구권 최근 반영 내용 알려줘",
            "규제문서 근거 + 페이지/파일 Citation + 정책 해석",
        ),
        (
            "심사결과 이미지 상담",
            "심사결과 이미지 업로드",
            "OCR 추출 + 거절사유/한도근거/보완항목 요약",
        ),
        (
            "상품 아이디어 생성",
            "당사 상품 적용 원리금 계산기 만들어줘",
            "부족정보 역질문 + 상품정보 기반 제안안",
        ),
    ]
    for use_case, user_input, result_output in rows:
        cells = table.add_row().cells
        cells[0].text = use_case
        cells[1].text = user_input
        cells[2].text = result_output

    add_heading(doc, "15. 기대효과", level=1)
    add_para(doc, "정량적 효과(목표)")
    add_bullets(
        doc,
        [
            "상담 응대시간 30~50% 단축",
            "규제근거 탐색시간 60% 이상 단축",
            "거절사유 설명 정확도/일관성 20%p 개선",
            "상품기획 아이디어 도출 리드타임 40% 단축",
        ],
    )
    add_para(doc, "정성적 효과")
    add_bullets(
        doc,
        [
            "상담원 설명 신뢰도 향상(근거 기반 응대)",
            "심사·기획·영업·IT 공통 언어(온톨로지) 확보",
            "정책 변경 대응 민첩성 강화",
            "데이터 기반 의사결정 문화 정착",
        ],
    )

    output_path = r"C:\work\ai-log-agent\AI_에이전트_대회_결과보고서.docx"
    doc.save(output_path)
    print(output_path)


if __name__ == "__main__":
    main()
